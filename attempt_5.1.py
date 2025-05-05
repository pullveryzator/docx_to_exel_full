from docx import Document
import pandas as pd
import os
import re
from time import sleep
from openpyxl import load_workbook

def parse_docx_to_excel(input_file, output_file):
    doc = Document(input_file)
    data = []
    skip_phrases = {'Сложение и вычитание натуральных чисел', 
                   'Натуральные числа',
                   'Умножение и деление натуральных чисел',
                   'Задачи «на части»',
                   'Задачи на нахождение двух чисел по их сумме и разности',
                   'Задачи на движение по реке',
                   'Задачи на движение',
                   'Разные задачи',
                   'Дроби',
                   'Вводные задачи',
                   'Сложение и вычитание обыкновенных дробей',
                   'Умножение и деление обыкновенных дробей',
                   'Задачи «на бассейны» и другие',
                   'Пропорции',
                   'Задачи на прямую и обратную пропорциональность',
                   'Проценты',
                   'Нахождение процентов числа',
                   'Нахождение процентного отношения',
                   'Сложные задачи на проценты',
                   'Уравнения',
                   'Решение задач с помощью уравнений',
                   'Более сложные задачи, решаемые уравнением',
                   'Задачи на повторение',
                   'Нахождение части числа и числа по его части',
                   'Нахождение числа по его процентам',
        }
    
    for para in doc.paragraphs:
        text = para.text
        if "Ответы и советы" in text:
            break
        if any(skip_phrase in text for skip_phrase in skip_phrases):
            continue
        if text:
            # Разделяем ID и задачу по первому вхождению табуляции или точки с пробелом
            if '\t' in text:
                parts = text.split('\t', 1)
            else:
                continue

            # sleep(0.5)

            id_part = parts[0]
            task_part = parts[1]
            if '.' in id_part:
                main_num = id_part
                subtask_parts = task_part.split('\t', 1)
                print(main_num, subtask_parts[0])
                if len(subtask_parts) == 1:
                    if '*' in main_num:
                        main_num = main_num.replace('*', '')
                        subtask_parts[0] = '*' + subtask_parts[0]
                    data.append({
                        'id_tasks_book': main_num,
                        'task': subtask_parts[0],
                        'answers': 'Отсутствует',
                        'paragraph': 1,
                        'classes': '5;6',
                        'topic_id': 1,
                        'level': 1
                    })
                if len(subtask_parts) == 2:
                    slave_num = subtask_parts[0].replace(')', '')
                    print(main_num + slave_num, subtask_parts[1])
                    if '*' in slave_num:
                        slave_num = slave_num.replace('*', '')
                        subtask_parts[1] = '*' + subtask_parts[1]
                    data.append({
                    'id_tasks_book': main_num + slave_num,
                    'task': subtask_parts[1],
                    'answers': 'Отсутствует',
                    'paragraph': 1,
                    'classes': '5;6',
                    'topic_id': 1,
                    'level': 1
                })
                    
            slave_num = id_part.replace(')', '')
            if main_num.strip() != slave_num.strip():
                print(main_num + slave_num, task_part)
                if '*' in slave_num:
                    slave_num = slave_num.replace('*', '')
                    task_part = '*' + task_part
                data.append({
                    'id_tasks_book': main_num + slave_num,
                    'task': task_part,
                    'answers': 'Отсутствует',
                    'paragraph': 1,
                    'classes': '5;6',
                    'topic_id': 1,
                    'level': 1
                })

    df = pd.DataFrame(data)
    with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='tasks', index=False)

def parse_toc_to_excel(docx_path, output_file):
    doc = Document(docx_path)
    sections = []
    last_main_section_id = 0
    found_toc = False

    section_pattern = re.compile(r'^(\d+)\.\s+(.*?)\s*\d*$')
    subsection_pattern = re.compile(r'^(\d+\.\d+)\.\s+(.*?)\s*\d*$')
    
    for para in doc.paragraphs:
        text = para.text.strip()

        if not found_toc:
            if text.lower() == "оглавление":
                found_toc = True
            continue
        
        if found_toc and not text:
            break

        main_section_match = section_pattern.match(text)
        subsection_match = subsection_pattern.match(text)
        
        if main_section_match and not '.' in main_section_match.group(1):
            section_num = main_section_match.group(1)
            section_name = main_section_match.group(2)
            
            sections.append({
                'id': len(sections) + 1,
                'name': f"{section_num}.{section_name}",
                'parent': 0
            })
            last_main_section_id = len(sections)
            
        elif subsection_match:
            subsection_num = subsection_match.group(1)
            subsection_name = subsection_match.group(2)
            
            sections.append({
                'id': len(sections) + 1,
                'name': f"{subsection_num}.{subsection_name}",
                'parent': last_main_section_id
            })
    
    if not sections:
        print("Оглавление не найдено в документе!")
        return
    
    df = pd.DataFrame(sections)
    
    if os.path.exists(output_file):
        book = load_workbook(output_file)
        if 'table_of_contents' in book.sheetnames:
            book.remove(book['table_of_contents'])

        with pd.ExcelWriter(output_file, engine='openpyxl', mode='a') as writer:
            df.to_excel(writer, sheet_name='table_of_contents', index=False)
    else:
        df.to_excel(output_file, sheet_name='table_of_contents', index=False)
    
    print(f"Оглавление сохранено в {output_file}")

if __name__ == "__main__":
    docx_path = "tekstovye_zadachi_po_matematike.docx"
    output_file="tasks.xlsx"
    parse_docx_to_excel(docx_path, output_file)
    parse_toc_to_excel(docx_path, output_file)

