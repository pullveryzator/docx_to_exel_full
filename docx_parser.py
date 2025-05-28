import re
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document

from ai_solution import add_ai_solution_to_excel
from classifier import process_topics
from constants import (ANSWER_COLUMN, AUTHOR_DATA, AUTHOR_SHEET_NAME, CLASSES,
                       CLASSES_COLUMN, DOCX_PATH, ID_TASK_COLUMN, OUTPUT_FILE,
                       PARAGRAPH_COLUMN, TASK_COLUMN, TASK_SHEET_NAME,
                       TOC_SHEET_NAME, TRIM_CHARS)
from decorators import validate_docx_file, validate_excel_file
from fixes import (fix_degree_to_star, fix_difficult_tasks_symb,
                   fix_trailing_dots)
from utils import (excel_to_dict, find_matching_paragraph, is_main_task,
                   is_subtask, reorder_sheets, save_to_excel)


@validate_docx_file
def parse_toc_to_excel(input_file:str, output_file:str):
    """Парсинг оглавления в Excel."""
    doc = Document(input_file)
    sections = []
    last_main_section_id = 0
    found_toc = False

    section_pattern = re.compile(r'^(\d+)\.\s+(.*?)\s*\d*$')
    subsection_pattern = re.compile(r'^(\d+\.\d+)\.\s+(.*?)\s*\d*$')
    
    for para in doc.paragraphs:
        text = para.text.strip()

        if not found_toc:
            if text.lower() == "оглавление":
                found_toc = True
            continue
        
        if found_toc and not text:
            break

        main_section_match = section_pattern.match(text)
        subsection_match = subsection_pattern.match(text)
        
        if main_section_match and not '.' in main_section_match.group(1):
            section_num = main_section_match.group(1)
            section_name = main_section_match.group(2)
            
            sections.append({
                'id': len(sections) + 1,
                'name': f"{section_num}.{section_name}",
                'parent': 0
            })
            last_main_section_id = len(sections)
            
        elif subsection_match:
            subsection_num = subsection_match.group(1)
            subsection_name = subsection_match.group(2)
            full_name = f"{subsection_num}.{subsection_name}"
            full_name = fix_trailing_dots(full_name)
            
            sections.append({
                'id': len(sections) + 1,
                'name': full_name,
                'parent': last_main_section_id
            })

    save_to_excel(data=sections, output_file=output_file, sheet_name=TOC_SHEET_NAME)


@validate_docx_file
def parse_docx_to_excel(input_file:str, output_file:str):
    """Парсинг текста задач в Excel."""
    doc = Document(input_file)
    data = []
    toc = excel_to_dict(output_file)

    for para in doc.paragraphs:
        text = para.text.strip()
        text = fix_degree_to_star(text)
        if "Ответы и советы" in text:
            break

        cleaned_text = re.sub(r'(\d+\.\d*\.?)\s+', r'\1', text)
        new_paragraph_id = find_matching_paragraph(cleaned_text, toc, trim_chars=TRIM_CHARS)

        if new_paragraph_id:
            paragraph_id = new_paragraph_id
            print(f"Найдено совпадение c оглавлением: {cleaned_text}")
            continue

        if '\t' in text:
            parts = text.split('\t', 1)
        id_part = parts[0].strip()
        task_part = parts[1].strip()
        if '.' in id_part:
            main_num = id_part
            subtask_parts = task_part.split('\t', 1)
            if len(subtask_parts) == 1:
                main_num, subtask_parts = fix_difficult_tasks_symb(main_num, subtask_parts, 0)
                data.append({
                    ID_TASK_COLUMN: main_num,
                    TASK_COLUMN: subtask_parts[0],
                    ANSWER_COLUMN: '',
                    PARAGRAPH_COLUMN: paragraph_id,
                    CLASSES_COLUMN: CLASSES,
                })
            if len(subtask_parts) == 2:
                slave_num = subtask_parts[0].replace(')', '')
                slave_num, subtask_parts = fix_difficult_tasks_symb(slave_num, subtask_parts, 1)
                data.append({
                ID_TASK_COLUMN: main_num + slave_num,
                TASK_COLUMN: subtask_parts[1],
                ANSWER_COLUMN: '',
                PARAGRAPH_COLUMN: paragraph_id,
                CLASSES_COLUMN: CLASSES,
            })
                
        slave_num = id_part.replace(')', '')
        if main_num.strip() != slave_num.strip():
            slave_num, task_part = fix_difficult_tasks_symb(slave_num, task_part)
            data.append({
                ID_TASK_COLUMN: main_num + slave_num,
                TASK_COLUMN: task_part,
                ANSWER_COLUMN: '',
                PARAGRAPH_COLUMN: paragraph_id,
                CLASSES_COLUMN: CLASSES,
            })

    save_to_excel(data=data, output_file=output_file, sheet_name=TASK_SHEET_NAME)


@validate_docx_file
def parse_answers(docx_path: str, output_file: str):
    """Парсинг ответов в Excel."""
    doc = Document(docx_path)
    answers_dict = {}
    answer_block_re = re.compile(r'(\d+)\.(.*?)(?=\d+\.|\Z)', re.DOTALL)
    answer_item_re = re.compile(
        r'([а-я]\)|\d+\))?\s*([^;.]*[;.]?)',
        re.DOTALL
    )

    full_text = "\n".join([para.text for para in doc.paragraphs])

    answers_start = full_text.find("Ответы и советы")
    answers_end = full_text.find('Оглавление')
    answers_text = full_text[answers_start:answers_end]

    for block in answer_block_re.finditer(answers_text):
        main_num = block.group(1)
        content = block.group(2).strip()

        answer_items = [a.strip() for a in content.split(';') if a.strip()]
        
        for item in answer_items:
            match = answer_item_re.match(item)
            if not match:
                continue
                
            subtask = match.group(1)
            answer_text = match.group(2).strip()

            if subtask:
                subtask = subtask.replace(')', '')
                task_id = f"{main_num}.{subtask}" if subtask.isalpha() else f"{main_num}.{subtask}"
            else:
                task_id = f"{main_num}."

            answers_dict[task_id] = re.sub(r'[.,;]$', '', answer_text).strip()

    tasks_df = pd.read_excel(output_file, sheet_name=TASK_SHEET_NAME)
    tasks_df[ANSWER_COLUMN] = tasks_df[ID_TASK_COLUMN].map(answers_dict).fillna('Отсутствует')
    with pd.ExcelWriter(output_file, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
        tasks_df.to_excel(writer, index=False, sheet_name=TASK_SHEET_NAME)

@validate_excel_file
def add_author(output_file:str, author_data: list[dict]):
    """Добавление к Excel файлу листа авторов."""
    save_to_excel(data=author_data, output_file=output_file, sheet_name=AUTHOR_SHEET_NAME)


@validate_excel_file
def process_composite_tasks(output_file: str) -> None:
    """Обрабатывает составные задачи в Excel-файле, объединяя условия основных задач с подзадачами.

    Функция выполняет следующие действия:
    1. Находит все основные задачи (формата "X.") с подзадачами
    2. Добавляет условие основной задачи в начало условия каждой подзадачи
    3. Удаляет обработанные основные задачи из файла
    4. Сохраняет модифицированные данные обратно в исходный файл

    Args:
        output_file: Путь к Excel-файлу с задачами. Должен содержать столбцы:
                    - 'id_tasks_book' - идентификаторы задач
                    - 'task' - тексты условий задач

    Returns:
        None

    Examples:
        Если в файле есть задачи:
        id_tasks_book | task
        '5.'          | 'Найти сумму'
        '5.1'         | 'чисел 2 и 3'
        
        После обработки:
        id_tasks_book | task
        '5.1'         | 'Найти сумму чисел 2 и 3'
    """
    df: pd.DataFrame = pd.read_excel(output_file, sheet_name=TASK_SHEET_NAME)

    modified_df: pd.DataFrame = df.copy()

    main_tasks: List[str] = []
    tasks_hierarchy: Dict[str, Dict[str, Any]] = {}
    
    for _, row in df[df[ID_TASK_COLUMN].apply(is_main_task)].iterrows():
        main_num: str = str(row[ID_TASK_COLUMN]).rstrip('.')
        subtasks: List[str] = df[df[ID_TASK_COLUMN].apply(lambda x: is_subtask(x, main_num))][ID_TASK_COLUMN].tolist()
        
        if subtasks:
            main_tasks.append(row[ID_TASK_COLUMN])
            tasks_hierarchy[row[ID_TASK_COLUMN]] = {
                'subtasks': subtasks,
                'main_condition': row[TASK_COLUMN]
            }

    for main_task, data in tasks_hierarchy.items():
        main_condition: str = data['main_condition']

        subtask_indices = df[df[ID_TASK_COLUMN].isin(data['subtasks'])].index

        for idx in subtask_indices:
            original_task: Optional[str] = modified_df.at[idx, TASK_COLUMN]
            if pd.notna(original_task):
                modified_df.at[idx, TASK_COLUMN] = f"{main_condition} {original_task}"
            else:
                modified_df.at[idx, TASK_COLUMN] = main_condition

    modified_df = modified_df[~modified_df[ID_TASK_COLUMN].isin(main_tasks)]

    with pd.ExcelWriter(output_file, engine='openpyxl', mode='w') as writer:
        modified_df.to_excel(writer, sheet_name=TASK_SHEET_NAME, index=False)
    
    if tasks_hierarchy:
        print("\nБыли обработаны следующие основные задачи и их подзадачи:")
        for main_task, data in tasks_hierarchy.items():
            print(f"\nОсновная задача: {main_task}")
            print(f"Условие: {data['main_condition']}")
            print(f"Подзадачи: {', '.join(data['subtasks'])}")
        print(f"\nОбработано {len(tasks_hierarchy)} составных задач.")
    else:
        print("Подходящих под условие обработки задач не нашлось.")

if __name__ == "__main__":

    parse_toc_to_excel(DOCX_PATH, OUTPUT_FILE)
    parse_docx_to_excel(DOCX_PATH, OUTPUT_FILE)
    add_author(OUTPUT_FILE, AUTHOR_DATA)
    parse_answers(DOCX_PATH, OUTPUT_FILE)
    process_composite_tasks(OUTPUT_FILE)
    add_ai_solution_to_excel(OUTPUT_FILE)
    reorder_sheets(OUTPUT_FILE)
    process_topics(OUTPUT_FILE)
    